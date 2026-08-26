from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

from generate_etral_technical_doc import (
    ASSETS,
    SCREENSHOT_DIR,
    NAVY,
    BLUE,
    MUTED,
    INK,
    configure_doc,
    add_table,
    add_paragraph,
    add_bullets,
    add_code_block,
    add_figure,
    add_screenshot,
    set_run_font,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "ETRAL_guia_explicacion_codigo_y_construccion.docx"


def build_doc():
    doc = Document()
    configure_doc(doc)

    # Portada orientada a la sustentacion.
    doc.add_paragraph().paragraph_format.space_after = Pt(80)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ETRAL")
    set_run_font(r, size=28, color=BLUE, bold=True)
    p.paragraph_format.space_after = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GUIA PARA EXPLICAR LA CONSTRUCCION\nDEL SISTEMA WEB Y GEMELO DIGITAL")
    set_run_font(r, size=21, color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Base del proyecto, lectura del codigo y guion de sustentacion")
    set_run_font(r, size=12.5, color=MUTED, italic=True)
    p.paragraph_format.space_after = Pt(45)
    add_table(doc, ["Documento", "Proposito"], [
        ("Guia de explicacion", "Ayudar a comentar como se construyo la aplicacion."),
        ("Audiencia", "Tesistas, asesor y jurado de Ingenieria Industrial."),
        ("Base tecnica", "React, Supabase/PostgreSQL, FastAPI y motor de simulacion."),
    ], [2700, 6660])
    doc.add_paragraph().paragraph_format.space_after = Pt(70)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Documento complementario al anexo tecnico de arquitectura")
    set_run_font(r, size=10, color=MUTED)
    doc.add_page_break()

    doc.add_heading("Como usar esta guia", level=1)
    add_paragraph(doc, "Este documento no busca que se memorice cada linea de codigo. Su objetivo es que se pueda explicar la razon de cada modulo, el recorrido de los datos y las decisiones que hacen que la aplicacion sea util para una planta industrial.")
    add_paragraph(doc, "La explicacion debe ir de lo general a lo particular: primero el problema industrial, luego la solucion, despues la arquitectura y finalmente el codigo que implementa cada responsabilidad.")

    doc.add_heading("1. Base del proyecto que deben explicar", level=1)
    doc.add_heading("1.1 Problema que origina la aplicacion", level=2)
    add_paragraph(doc, "La operacion de planta necesita controlar productos, rutas, fases, actividades, materiales, ordenes CECO, personal, equipos e incidencias. Cuando estos datos estan dispersos o se actualizan sin trazabilidad, es dificil saber que orden esta retrasada, que material limita el plan y que capacidad real queda disponible.")
    add_paragraph(doc, "La aplicacion responde a ese problema centralizando la informacion operativa y agregando un gemelo digital capaz de evaluar escenarios antes de tomar una decision en planta.")
    doc.add_heading("1.2 Objetivo de la solucion", level=2)
    add_paragraph(doc, "Construir un sistema web que registre y consulte la operacion de produccion, y que permita simular cambios de demanda, inventario, turnos, disponibilidad laboral y prioridades para anticipar restricciones y apoyar la toma de decisiones.")
    doc.add_heading("1.3 Alcance", level=2)
    add_bullets(doc, [
        "Registrar maestros de productos, fases, actividades, materiales y recursos.",
        "Definir rutas de fabricacion y listas de materiales BOM.",
        "Registrar ordenes CECO, avance, WIP, movimientos, calidad e incidencias.",
        "Calcular indicadores de capacidad, atrasos, materiales en riesgo y cumplimiento.",
        "Comparar un plan operativo con escenarios what-if sin modificar la realidad registrada.",
    ])
    doc.add_heading("1.4 Lo que el sistema no pretende hacer", level=2)
    add_bullets(doc, [
        "No reemplaza la decision del responsable de planta.",
        "No ejecuta automaticamente un cambio en las ordenes reales al simular.",
        "No es un ERP completo ni un sistema contable.",
        "No inventa datos: necesita maestros, avances y registros de operacion para calibrarse.",
    ])

    doc.add_heading("2. Orden recomendado para explicar la construccion", level=1)
    add_table(doc, ["Paso", "Que se explica", "Pregunta que responde"], [
        ("1", "Proceso industrial y problema", "Por que era necesario construirlo?"),
        ("2", "Requisitos y alcance", "Que registra y que simula?"),
        ("3", "Modelo de datos", "Como se representa una orden y sus materiales?"),
        ("4", "Arquitectura", "Como se separan interfaz, servicios y persistencia?"),
        ("5", "Implementacion", "Que archivo cumple cada responsabilidad?"),
        ("6", "Gemelo digital", "Como se construye el escenario y se calculan resultados?"),
        ("7", "Demostracion y evidencias", "Como se comprueba que funciona?"),
    ], [900, 4150, 4310])
    add_paragraph(doc, "Frase de transicion recomendada: Primero definimos el proceso y los datos que necesitaba la planta; despues construimos las capas de la aplicacion para que el registro operativo y el analisis de escenarios pudieran evolucionar sin mezclarse.")

    doc.add_heading("3. Mapa de la base de codigo", level=1)
    add_table(doc, ["Ruta", "Responsabilidad", "Como comentarla"], [
        ("src/App.jsx", "Composicion de pantallas, estado y formularios.", "Es el punto de ensamblaje de la experiencia web."),
        ("src/services/repository.js", "Seleccion del repositorio de datos.", "La vista trabaja con un contrato, no con SQL."),
        ("src/services/supabaseRepository.js", "Lectura y escritura en Supabase.", "Convierte acciones de usuario en operaciones de persistencia."),
        ("src/services/localRepository.js", "Datos de demostracion o mock.", "Permite probar la interfaz sin depender de la base real."),
        ("src/services/twinApi.js", "Adaptacion al motor del gemelo.", "Elige JavaScript o FastAPI sin cambiar la vista."),
        ("src/lib/simulator.js", "Calculo local de escenarios y KPIs.", "Contiene reglas de capacidad, materiales, atrasos y resultados."),
        ("backend/app/main.py", "Endpoints HTTP del motor Python.", "Recibe estructuras validadas y delega al dominio."),
        ("src/supabase/schema.sql", "Tablas, claves y restricciones.", "Es la base persistente del modelo de planta."),
    ], [2600, 3350, 3410])
    add_paragraph(doc, "Cuando el jurado pregunte donde esta una regla, primero se identifica si es una regla de interfaz, persistencia o dominio. Esa clasificacion evita responder con un archivo incorrecto.")

    doc.add_heading("4. Como comentar el codigo", level=1)
    doc.add_heading("4.1 Punto de entrada: composicion de la interfaz", level=2)
    add_paragraph(doc, "App.jsx no debe explicarse como un archivo gigante, sino como el punto donde se componen las pantallas y se conectan acciones del usuario con servicios. La idea es mostrar que la interfaz representa el proceso: inicio, produccion, productos, fases, inventario, recursos y simulacion.")
    add_bullets(doc, [
        "El usuario inicia una accion desde una vista.",
        "La vista valida y prepara el formulario.",
        "El servicio ejecuta la operacion.",
        "El repositorio actualiza los datos.",
        "La interfaz refresca el dataset y muestra el nuevo estado.",
    ])

    doc.add_heading("4.2 Repository: por que no se consulta la base desde cada vista", level=2)
    add_paragraph(doc, "Este fragmento muestra una decision de desacoplamiento. La interfaz solicita un repositorio y el entorno define si se usa modo mock o Supabase.")
    add_code_block(doc, '''import { localRepository } from "./localRepository.js";
import { hasSupabaseConfig, supabaseRepository } from "./supabaseRepository.js";

export function getRepository() {
  if (import.meta.env.VITE_DATA_MODE === "mock") return localRepository;
  return hasSupabaseConfig() ? supabaseRepository : supabaseRepository;
}''')
    add_paragraph(doc, "Que decir: El patron Repository concentra el acceso a datos. Asi, una pantalla no necesita saber si los datos vienen de Supabase o de un conjunto de demostracion. La prueba de la interfaz se vuelve mas sencilla y se reduce el acoplamiento.")
    add_paragraph(doc, "Pregunta posible: Por que existe localRepository? Respuesta: Permite desarrollar y demostrar la interfaz con datos controlados cuando no se desea depender de una conexion o de datos productivos.")

    doc.add_heading("4.3 SupabaseRepository: como se carga el dataset", level=2)
    add_paragraph(doc, "El repositorio consulta las tablas requeridas y opcionales en paralelo. Las tablas requeridas detienen la carga si fallan; las opcionales permiten que el sistema siga funcionando si una capacidad adicional aun no esta configurada.")
    add_code_block(doc, '''const requiredTables = ["flow_stages", "stage_activities", "body_types",
  "product_routes", "inventory_items", "bom_items", "ceco_orders"];
const optionalTables = ["work_shifts", "personnel", "equipment",
  "work_calendar", "resource_assignments", "operational_incidents"];
const tables = [...requiredTables, ...optionalTables];
const results = await Promise.all(
  tables.map((table) => supabase.from(table).select("*"))
);''')
    add_paragraph(doc, "Que decir: La aplicacion necesita un dataset coherente. Por eso se cargan en conjunto las entidades que alimentan los tableros y el gemelo. El repositorio tambien centraliza el manejo de errores para no repetirlo en cada componente.")

    doc.add_heading("4.4 Operaciones de escritura y trazabilidad", level=2)
    add_paragraph(doc, "Registrar un movimiento no consiste solo en cambiar el saldo de un material. Tambien se crea una fila en inventory_movements, de modo que pueda reconstruirse que ocurrio, con que cantidad y para que CECO.")
    add_code_block(doc, '''const { error } = await supabase
  .from("inventory_movements")
  .insert({
    id: `mov-${Date.now()}`,
    type: payload.type,
    code: payload.code,
    ceco: payload.ceco || "",
    quantity,
    note: payload.note,
  });
if (error) throw error;''')
    add_paragraph(doc, "Que decir: La trazabilidad es parte del objetivo industrial. El sistema conserva el hecho operativo y no solo el valor final del inventario.")

    doc.add_page_break()
    doc.add_heading("4.5 Snapshot y seleccion del motor", level=2)
    add_paragraph(doc, "Antes de simular, twinApi.js transforma el dataset en un snapshot con materiales, ordenes, rutas y recursos. Luego elige el motor local o la API Python.")
    add_code_block(doc, '''export async function runTwinSimulation(dataset, draft) {
  if (twinEngine !== "python") {
    return runDigitalTwin(dataset, draft);
  }

  const response = await fetch(`${baseUrl}/api/v1/simulations`, {
    method: "POST",
    headers: { "Content-Type": "application/json" },
    body: JSON.stringify({
      name: "Simulacion desde interfaz",
      input: { snapshot: snapshotFromDataset(dataset) },
    }),
  });
  return response.json();
}''')
    add_paragraph(doc, "Que decir: El adaptador evita que la interfaz conozca la implementacion del motor. El snapshot protege la informacion real porque la corrida calcula sobre una fotografia del estado operativo.")

    doc.add_heading("4.6 Motor del gemelo digital", level=2)
    add_code_block(doc, '''export function runDigitalTwin(dataset, params = {}) {
  const calibration = params.calibrationData ?? calibrateDigitalTwin(dataset);
  const normalized = {
    horizonDays: Number(params.horizonDays ?? 14),
    laborAvailability: Number(params.laborAvailability ?? 1),
    shiftsPerDay: Number(params.shiftsPerDay ?? 1),
  };
  // Calcula capacidad, restricciones, atrasos y KPIs.
}''')
    add_paragraph(doc, "Que decir: El motor normaliza los parametros, calibra el modelo y calcula resultados comparables. Entre sus salidas se encuentran capacidad por fase, materiales criticos, ordenes retrasadas, throughput, lead time y cumplimiento del plan.")
    add_paragraph(doc, "Importante: No se debe afirmar que la simulacion predice el futuro con certeza. Se debe explicar que evalua escenarios bajo supuestos y datos disponibles, entregando informacion para decidir.")

    doc.add_heading("4.7 API Python y separacion de dominio", level=2)
    add_code_block(doc, '''app = FastAPI(title="ETRAL Digital Twin API", version="0.1.0")

@app.post("/api/v1/simulations")
def run_simulation(run: SimulationRun) -> dict:
    return {"name": run.name,
            "result": simulate_comparison(run.input)}''')
    add_paragraph(doc, "Que decir: El endpoint recibe una estructura validada y delega el trabajo a simulate_comparison. La logica de negocio queda separada del transporte HTTP y puede probarse de manera independiente.")

    doc.add_heading("5. Como explicar la base de datos", level=1)
    add_paragraph(doc, "La base de datos representa el proceso industrial en cuatro grupos: maestros, operacion, recursos y trazabilidad. El producto define una ruta y un BOM; la orden CECO recorre la ruta; el inventario abastece el BOM; y los recursos permiten estimar capacidad.")
    add_table(doc, ["Grupo", "Tablas principales", "Ejemplo de pregunta"], [
        ("Maestros", "body_types, flow_stages, stage_activities, inventory_items", "Que se fabrica y como se fabrica?"),
        ("Operacion", "ceco_orders, stage_inventory, ceco_activity_progress", "En que fase esta cada CECO?"),
        ("Recursos", "personnel, equipment, work_calendar, assignments", "Con que capacidad se cuenta?"),
        ("Trazabilidad", "inventory_movements, operation_logs, quality_checks, incidents", "Que ocurrio y cuando?"),
        ("Analitica", "simulation_runs", "Que escenario se evaluo y con que resultado?"),
    ], [1700, 4300, 3360])
    add_paragraph(doc, "La regla de oro para explicarla es: cada dato operativo debe tener una entidad donde vivir y una relacion que permita rastrearlo hasta el producto, la orden, la fase o el recurso involucrado.")

    doc.add_heading("6. Como explicar el gemelo digital", level=1)
    add_table(doc, ["Elemento", "Explicacion para el jurado"], [
        ("Entrada", "Snapshot del dataset y parametros del escenario."),
        ("Modelo", "Capacidad por fase, tiempos estandar, inventario, rutas, recursos e incertidumbre."),
        ("Calculo", "Distribuye capacidad, verifica restricciones y compara el escenario con el plan."),
        ("Salida", "KPIs, atrasos, cuellos de botella, stockouts, alertas y nivel de cumplimiento."),
        ("Seguridad operacional", "El escenario no modifica los registros reales de la planta."),
    ], [2100, 7260])
    add_paragraph(doc, "La diferencia entre un tablero y el gemelo es que el tablero describe el estado actual, mientras que el gemelo permite explorar que pasaria si cambia una condicion del sistema.")

    doc.add_heading("7. Demostracion recomendada", level=1)
    add_paragraph(doc, "La demostracion debe ser corta y con un hilo unico. Se recomienda usar una orden CECO, mostrar su producto y ruta, revisar un material, consultar recursos y finalizar con una simulacion.")
    add_bullets(doc, [
        "Inicio: explicar la situacion de la planta y los indicadores de alerta.",
        "Produccion: mostrar una orden CECO ubicada en una fase.",
        "Productos: mostrar que la plantilla contiene ruta y BOM.",
        "Inventario: explicar fisico, comprometido, disponible y proyeccion.",
        "Recursos: mostrar personal, horas efectivas, equipos e incidencias.",
        "Simulacion: modificar una variable y comparar el resultado sin alterar la operacion.",
    ])

    screenshots = [
        ("01_inicio.png", "Figura 1. Inicio: resumen de la situacion operativa."),
        ("02_produccion.png", "Figura 2. Produccion: tablero de ordenes CECO por fase."),
        ("03_productos_rutas.png", "Figura 3. Productos: ruta y estructura BOM."),
        ("05_inventario_bom.png", "Figura 4. Inventario: cobertura y proyeccion de materiales."),
        ("06_recursos.png", "Figura 5. Recursos: alimentacion operativa del gemelo."),
        ("07_simulacion.png", "Figura 6. Simulacion: configuracion del escenario."),
    ]
    available = [(SCREENSHOT_DIR / name, caption) for name, caption in screenshots if (SCREENSHOT_DIR / name).exists()]
    for index, (path, caption) in enumerate(available):
        if index == 0:
            doc.add_page_break()
            doc.add_heading("8. Evidencia visual para comentar", level=1)
            add_paragraph(doc, "Estas capturas sirven como apoyo durante la explicacion. No se deben leer como una lista de botones: cada interfaz debe conectarse con una necesidad del proceso industrial.")
        else:
            doc.add_page_break()
        add_screenshot(doc, path, caption)

    doc.add_page_break()
    doc.add_heading("9. Guion breve de sustentacion", level=1)
    add_paragraph(doc, "El siguiente guion puede adaptarse al estilo de cada integrante del equipo:")
    add_table(doc, ["Momento", "Texto sugerido"], [
        ("Problema", "La planta necesitaba centralizar la trazabilidad de ordenes, materiales, fases y recursos."),
        ("Solucion", "Construimos una aplicacion web para registrar la operacion y un gemelo digital para evaluar escenarios."),
        ("Arquitectura", "Separamos interfaz, servicios, persistencia y dominio para que cada parte tenga una responsabilidad clara."),
        ("Codigo", "El Repository desacopla la vista de Supabase; twinApi desacopla la vista del motor; simulator concentra el calculo."),
        ("Base de datos", "Las relaciones conectan productos, rutas, BOM, ordenes, inventario, recursos y trazabilidad."),
        ("Resultado", "El sistema permite anticipar restricciones y apoyar decisiones sin modificar la operacion real al simular."),
    ], [1800, 7560])

    doc.add_heading("10. Preguntas probables del jurado", level=1)
    add_table(doc, ["Pregunta", "Respuesta recomendada"], [
        ("Por que React?", "Permite construir una interfaz modular y actualizar el estado de la operacion sin recargar toda la pagina."),
        ("Por que Supabase?", "Proporciona PostgreSQL y acceso estructurado desde el cliente, adecuado para el modelo relacional del proyecto."),
        ("Por que dos motores?", "El motor JavaScript facilita la respuesta inmediata; Python permite ampliar el analisis sin cambiar la interfaz."),
        ("El gemelo reemplaza al jefe de planta?", "No. Es una herramienta de evaluacion de escenarios; la decision sigue siendo humana y operativa."),
        ("Como se valida?", "Comparando resultados con datos historicos, tiempos estandar, avances registrados y escenarios controlados."),
        ("Que pasa si faltan datos?", "El resultado debe declararse como limitado y se debe priorizar la calidad de maestros y registros operativos."),
    ], [3500, 5860])

    doc.add_heading("11. Cierre y mejoras futuras", level=1)
    add_paragraph(doc, "La aplicacion queda preparada para ampliar autenticacion, roles, politicas RLS, mayor integracion con registros historicos y modelos de simulacion mas especializados. Estas mejoras no cambian la base: mantener trazabilidad, separar responsabilidades y proteger la operacion real durante la evaluacion de escenarios.")
    add_paragraph(doc, "La idea final que deben transmitir es sencilla: el sistema convierte datos dispersos de produccion en una representacion operativa consultable y en escenarios comparables para decidir mejor.")
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()
