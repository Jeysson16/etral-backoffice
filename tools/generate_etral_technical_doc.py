from pathlib import Path
from math import atan2, cos, sin, pi

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets" / "etral-diagrams"
SCREENSHOT_DIR = DOCS / "assets" / "etral-screenshots"
OUTPUT = DOCS / "ETRAL_documentacion_tecnica_sistema_web_y_gemelo_digital_con_capturas.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
TEAL = "287D8E"
GREEN = "2E7D5B"
GOLD = "A56A00"
RED = "A63A3A"
INK = "1F2937"
MUTED = "5B6673"
LIGHT = "EAF0F6"
LIGHT_TEAL = "E5F4F5"
LIGHT_GOLD = "FFF4D8"
WHITE = "FFFFFF"


def font(size, bold=False, mono=False):
    candidates = []
    if mono:
        candidates += ["C:/Windows/Fonts/consola.ttf", "C:/Windows/Fonts/cour.ttf"]
    else:
        candidates += [
            "C:/Windows/Fonts/calibri.ttf",
            "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/segoeui.ttf",
        ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


F_SMALL = font(24)
F_BODY = font(28)
F_LABEL = font(30)
F_TITLE = font(42)
F_BIG = font(50)


def pc(color):
    return color if not isinstance(color, str) or color.startswith("#") else f"#{color}"


def text_size(draw, text, fnt):
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def center_text(draw, text, rect, fnt, fill=INK):
    x1, y1, x2, y2 = rect
    width, height = text_size(draw, text, fnt)
    draw.text((x1 + (x2 - x1 - width) / 2, y1 + (y2 - y1 - height) / 2 - 3), text, font=fnt, fill=pc(fill))


def wrap(draw, text, fnt, max_width):
    words = text.split()
    lines, line = [], ""
    for word in words:
        candidate = word if not line else f"{line} {word}"
        if text_size(draw, candidate, fnt)[0] <= max_width:
            line = candidate
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    return lines


def wrapped_center(draw, text, rect, fnt, fill=INK, gap=6):
    x1, y1, x2, y2 = rect
    lines = []
    for paragraph in text.split("\n"):
        lines.extend(wrap(draw, paragraph, fnt, x2 - x1 - 26))
    line_h = text_size(draw, "Ag", fnt)[1] + gap
    total_h = len(lines) * line_h - gap
    current_y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        w, _ = text_size(draw, line, fnt)
        draw.text((x1 + (x2 - x1 - w) / 2, current_y), line, font=fnt, fill=pc(fill))
        current_y += line_h


def base_diagram(title, subtitle, width=1800, height=1120):
    image = Image.new("RGB", (width, height), pc(WHITE))
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, width, 112), fill=pc(NAVY))
    draw.text((52, 28), title, font=F_TITLE, fill=pc(WHITE))
    draw.text((54, 132), subtitle, font=F_SMALL, fill=pc(MUTED))
    return image, draw


def box(draw, rect, title, detail="", fill=LIGHT, border=BLUE, title_color=NAVY):
    x1, y1, x2, y2 = rect
    draw.rounded_rectangle(rect, radius=20, fill=pc(fill), outline=pc(border), width=4)
    title_rect = (x1 + 14, y1 + 12, x2 - 14, y1 + 58)
    wrapped_center(draw, title, title_rect, F_LABEL, fill=title_color)
    if detail:
        wrapped_center(draw, detail, (x1 + 18, y1 + 64, x2 - 18, y2 - 16), F_SMALL, fill=INK)


def arrow(draw, start, end, color=BLUE, width=5, label=None):
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=pc(color), width=width)
    angle = atan2(y2 - y1, x2 - x1)
    length = 20
    left = (x2 - length * cos(angle - pi / 7), y2 - length * sin(angle - pi / 7))
    right = (x2 - length * cos(angle + pi / 7), y2 - length * sin(angle + pi / 7))
    draw.polygon([(x2, y2), left, right], fill=pc(color))
    if label:
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        draw.rounded_rectangle((mx - 65, my - 20, mx + 65, my + 20), radius=8, fill=pc(WHITE))
        center_text(draw, label, (mx - 60, my - 16, mx + 60, my + 16), F_SMALL, color)


def diamond(draw, center, width, height, label, fill=LIGHT_GOLD, border=GOLD):
    cx, cy = center
    points = [(cx, cy - height / 2), (cx + width / 2, cy), (cx, cy + height / 2), (cx - width / 2, cy)]
    draw.polygon(points, fill=pc(fill), outline=pc(border))
    wrapped_center(draw, label, (cx - width / 2 + 22, cy - height / 2 + 22, cx + width / 2 - 22, cy + height / 2 - 22), F_SMALL, fill=INK)


def save(image, name):
    path = ASSETS / f"{name}.png"
    image.save(path, "PNG", optimize=True)
    return path


def create_architecture():
    image, draw = base_diagram("Arquitectura por capas", "Separacion de responsabilidades del sistema ETRAL", 1800, 1160)
    layers = [
        ("Presentacion", "React + Vite\nApp.jsx, vistas y formularios", "EAF0F6", BLUE),
        ("Aplicacion", "Repositorio, adaptador del gemelo y reglas de coordinacion", "E5F4F5", TEAL),
        ("Dominio", "Capacidad, MRP, alertas, KPIs y simulacion what-if", "FFF4D8", GOLD),
        ("Infraestructura", "Supabase/PostgreSQL y API FastAPI", "F5EAF1", "8B4B8A"),
    ]
    y = 220
    for index, (title, detail, fill, border) in enumerate(layers):
        box(draw, (270, y, 1530, y + 155), title, detail, fill=fill, border=border)
        if index < len(layers) - 1:
            arrow(draw, (900, y + 155), (900, y + 206), color=border)
        y += 210
    draw.rounded_rectangle((270, 1055, 1530, 1115), radius=14, fill=pc("F2F4F7"))
    center_text(draw, "El gemelo opera sobre snapshots y no modifica los registros operativos.", (285, 1060, 1515, 1112), F_SMALL, NAVY)
    return save(image, "01_arquitectura_por_capas")


def create_application_diagram():
    image, draw = base_diagram("Diagrama de aplicaciones", "Interaccion entre usuario, aplicacion web, persistencia y motores de simulacion", 2000, 1160)
    boxes = {
        "user": (65, 420, 300, 600),
        "web": (380, 365, 720, 655),
        "repo": (810, 240, 1130, 430),
        "twin": (810, 650, 1130, 840),
        "db": (1260, 230, 1600, 450),
        "js": (1260, 610, 1600, 800),
        "api": (1690, 610, 1940, 800),
    }
    box(draw, boxes["user"], "Usuario de planta", "Registra y consulta la operacion", fill="F2F4F7", border=MUTED)
    box(draw, boxes["web"], "Aplicacion web", "React + Vite\nTableros, formularios, reportes", fill=LIGHT, border=BLUE)
    box(draw, boxes["repo"], "Capa de repositorio", "getRepository()\nSupabase o modo mock", fill=LIGHT_TEAL, border=TEAL)
    box(draw, boxes["twin"], "Adaptador del gemelo", "twinApi.js\nSelecciona el motor", fill=LIGHT_TEAL, border=TEAL)
    box(draw, boxes["db"], "Supabase PostgreSQL", "Maestros, ordenes, inventario, recursos y trazabilidad", fill="F5EAF1", border="8B4B8A")
    box(draw, boxes["js"], "Motor JavaScript", "simulator.js\nEjecucion en navegador", fill=LIGHT_GOLD, border=GOLD)
    box(draw, boxes["api"], "API FastAPI", "MRP y simulacion\nMotor Python", fill=LIGHT_GOLD, border=GOLD)
    arrow(draw, (300, 510), (380, 510), label="uso")
    arrow(draw, (720, 430), (810, 335), label="datos")
    arrow(draw, (1130, 335), (1260, 335), label="CRUD")
    arrow(draw, (720, 585), (810, 745), label="escenario")
    arrow(draw, (1130, 710), (1260, 705), label="browser")
    arrow(draw, (1130, 785), (1690, 785), label="python")
    return save(image, "02_diagrama_aplicaciones")


def create_components():
    image, draw = base_diagram("Diagrama de componentes", "Organizacion funcional de la aplicacion React y sus servicios", 1900, 1280)
    box(draw, (700, 190, 1200, 350), "App.jsx", "Compone las vistas y controla la navegacion", fill=LIGHT, border=BLUE)
    components = [
        (80, "ProductionBoard", "ordenes y capacidad"),
        (380, "Producto y ruta", "ProductList + FlowDrawer"),
        (680, "StagesView", "fases y actividades"),
        (980, "Inventario y BOM", "materiales y estructura"),
        (1280, "ExecutionPanel", "avance de produccion"),
        (1580, "ResourcesView", "personal y equipos"),
    ]
    for x, title, detail in components:
        box(draw, (x, 470, x + 250, 650), title, detail, fill="F2F4F7", border=MUTED)
        arrow(draw, (950, 350), (x + 125, 470), color=BLUE)
    box(draw, (535, 800, 840, 980), "getRepository()", "Contrato de datos\nCRUD sin SQL en la vista", fill=LIGHT_TEAL, border=TEAL)
    box(draw, (1050, 800, 1360, 980), "twinApi.js", "Adaptador de motor\nJS o Python", fill=LIGHT_TEAL, border=TEAL)
    for x in [205, 505, 805, 1105, 1405, 1705]:
        arrow(draw, (x, 650), (685, 800), color=MUTED, width=3)
    arrow(draw, (950, 650), (1205, 800), color=BLUE)
    box(draw, (330, 1080, 780, 1215), "supabaseRepository / localRepository", "Fuentes de datos intercambiables", fill="F5EAF1", border="8B4B8A")
    box(draw, (1080, 1080, 1530, 1215), "simulator.js / FastAPI", "Calculo de escenarios y KPIs", fill=LIGHT_GOLD, border=GOLD)
    arrow(draw, (690, 980), (555, 1080), color=TEAL)
    arrow(draw, (1205, 980), (1305, 1080), color=TEAL)
    return save(image, "03_diagrama_componentes")


def create_ceco_flow():
    image, draw = base_diagram("Flujo de una orden CECO", "Registro, validacion de materiales, ejecucion y cierre", 1850, 1600)
    nodes = [
        ((680, 190, 1170, 300), "Registrar pedido", "Datos del cliente, producto y fecha"),
        ((680, 370, 1170, 480), "Crear orden CECO", "Codigo unico y prioridad"),
        ((680, 550, 1170, 660), "Asignar ruta y fase inicial", "Secuencia definida por el producto"),
    ]
    for rect, title, detail in nodes:
        box(draw, rect, title, detail)
    arrow(draw, (925, 300), (925, 370))
    arrow(draw, (925, 480), (925, 550))
    diamond(draw, (925, 760), 420, 160, "BOM e inventario suficientes?")
    arrow(draw, (925, 660), (925, 680))
    box(draw, (105, 900, 555, 1035), "Registrar alerta", "Compra, reprogramacion o sustitucion", fill="FBEAEA", border=RED)
    box(draw, (695, 900, 1155, 1035), "Reservar o emitir material", "Movimiento de inventario", fill=LIGHT_TEAL, border=TEAL)
    box(draw, (1295, 900, 1745, 1035), "Actualizar avance", "WIP, actividades y recursos", fill=LIGHT, border=BLUE)
    arrow(draw, (715, 810), (330, 900), color=RED, label="No")
    arrow(draw, (925, 840), (925, 900), color=GREEN, label="Si")
    arrow(draw, (1155, 965), (1295, 965), color=BLUE)
    diamond(draw, (1520, 1160), 360, 150, "Ruta finalizada?")
    arrow(draw, (1520, 1035), (1520, 1085))
    box(draw, (1300, 1330, 1740, 1465), "Control de calidad y cierre", "Orden terminada y evidencia final", fill="E8F4EA", border=GREEN)
    box(draw, (635, 1330, 1075, 1465), "Mover a siguiente fase", "Continuar la ruta del producto", fill=LIGHT, border=BLUE)
    arrow(draw, (1425, 1210), (855, 1330), color=BLUE, label="No")
    arrow(draw, (1520, 1235), (1520, 1330), color=GREEN, label="Si")
    return save(image, "04_flujo_ceco")


def create_twin_flow():
    image, draw = base_diagram("Flujo del gemelo digital", "El escenario se calcula sobre una fotografia del estado operativo", 1800, 1450)
    flow = [
        ("Definir escenario", "demanda, turnos, ausentismo, materiales"),
        ("Leer dataset", "ordenes, inventario, rutas y recursos"),
        ("Construir snapshot", "copia inmutable para la simulacion"),
    ]
    y = 190
    for title, detail in flow:
        box(draw, (600, y, 1200, y + 120), title, detail)
        if y > 190:
            arrow(draw, (900, y - 60), (900, y))
        y += 180
    diamond(draw, (900, 790), 410, 155, "Motor configurado?")
    arrow(draw, (900, 670), (900, 710))
    box(draw, (150, 930, 610, 1065), "Motor JavaScript", "runDigitalTwin() en el navegador", fill=LIGHT_GOLD, border=GOLD)
    box(draw, (1190, 930, 1650, 1065), "Motor Python", "POST a FastAPI y servicios de dominio", fill=LIGHT_GOLD, border=GOLD)
    arrow(draw, (745, 850), (380, 930), label="browser")
    arrow(draw, (1055, 850), (1420, 930), label="python")
    box(draw, (600, 1165, 1200, 1305), "Comparar resultados", "capacidad, restricciones, atrasos, KPIs y alertas", fill="E8F4EA", border=GREEN)
    arrow(draw, (380, 1065), (760, 1165), color=GREEN)
    arrow(draw, (1420, 1065), (1040, 1165), color=GREEN)
    draw.rounded_rectangle((280, 1350, 1520, 1415), radius=16, fill=pc("F2F4F7"))
    center_text(draw, "Salida: recomendacion para decidir. No se actualizan CECO, inventario ni avances reales.", (300, 1354, 1500, 1412), F_SMALL, NAVY)
    return save(image, "05_flujo_gemelo_digital")


def create_database_diagram():
    image, draw = base_diagram("Diagrama logico de base de datos", "Agrupacion de entidades por responsabilidad operativa", 1950, 1300)
    groups = [
        ((70, 250, 520, 1050), "MAESTROS", ["Categorias, unidades y marcas", "flow_stages", "stage_activities", "body_types", "work_shifts"]),
        ((570, 250, 1170, 1050), "PLANEAMIENTO Y OPERACION", ["product_routes", "bom_items", "inventory_items", "ceco_orders", "stage_inventory", "ceco_activity_progress"]),
        ((1220, 250, 1880, 1050), "RECURSOS, CONTROL Y ANALITICA", ["personnel y equipment", "work_calendar", "resource_assignments", "incidents y quality_checks", "movements y operation_logs", "simulation_runs"]),
    ]
    for rect, title, items in groups:
        x1, y1, x2, y2 = rect
        draw.rounded_rectangle(rect, radius=26, fill=pc("F8FAFC"), outline=pc(BLUE), width=4)
        draw.rounded_rectangle((x1, y1, x2, y1 + 85), radius=26, fill=pc(NAVY))
        center_text(draw, title, (x1 + 10, y1 + 12, x2 - 10, y1 + 74), F_LABEL, WHITE)
        yy = y1 + 120
        for item in items:
            box(draw, (x1 + 35, yy, x2 - 35, yy + 95), item, "", fill=LIGHT, border="B7C9DA")
            yy += 112
    arrow(draw, (520, 650), (570, 650), color=TEAL, label="define")
    arrow(draw, (1170, 650), (1220, 650), color=TEAL, label="ejecuta")
    draw.rounded_rectangle((70, 1120, 1880, 1215), radius=20, fill=pc(LIGHT_GOLD), outline=pc(GOLD), width=3)
    center_text(draw, "Los maestros definen el proceso; las tablas de operacion registran la realidad; analitica conserva la evidencia de las corridas.", (100, 1136, 1850, 1200), F_SMALL, NAVY)
    return save(image, "06_diagrama_base_datos")


def create_er_diagram():
    image, draw = base_diagram("Diagrama entidad-relacion", "Relaciones principales que sostienen la trazabilidad de produccion", 2200, 1500)
    nodes = {
        "producto": (90, 280, 420, 455, "body_types", "Producto fabricado"),
        "fases": (90, 690, 420, 865, "flow_stages", "Fases de produccion"),
        "ruta": (600, 270, 960, 455, "product_routes", "Ruta por producto y fase"),
        "bom": (600, 690, 960, 875, "bom_items", "Materiales requeridos"),
        "orden": (1160, 430, 1530, 630, "ceco_orders", "Orden de produccion"),
        "inventario": (600, 1090, 960, 1275, "inventory_items", "Existencia y stock comprometido"),
        "avance": (1740, 245, 2110, 430, "ceco_activity_progress", "Progreso de actividades"),
        "wip": (1740, 570, 2110, 755, "stage_inventory", "WIP por fase"),
        "asig": (1740, 895, 2110, 1080, "resource_assignments", "Personal, orden y actividad"),
        "recursos": (1160, 1030, 1530, 1215, "personnel / equipment", "Recursos disponibles"),
    }
    for rect in nodes.values():
        x1, y1, x2, y2, title, detail = rect
        box(draw, (x1, y1, x2, y2), title, detail, fill=LIGHT, border=BLUE)
    relations = [
        ((420, 365), (600, 365), "1:N"),
        ((420, 775), (600, 775), "1:N"),
        ((420, 365), (1160, 500), "1:N"),
        ((420, 775), (1160, 560), "1:N"),
        ((960, 365), (1160, 500), "ruta"),
        ((960, 780), (600, 1180), "material"),
        ((960, 780), (1160, 560), "BOM"),
        ((1530, 500), (1740, 335), "1:N"),
        ((1530, 560), (1740, 660), "1:N"),
        ((1530, 600), (1740, 985), "1:N"),
        ((1530, 1120), (1740, 985), "1:N"),
    ]
    for start, end, label in relations:
        arrow(draw, start, end, color=TEAL, width=4, label=label)
    draw.rounded_rectangle((90, 1340, 2110, 1435), radius=18, fill=pc("F2F4F7"))
    center_text(draw, "Una orden CECO vincula el producto y fase con avances, WIP y recursos; el BOM conecta producto, fase y materiales.", (120, 1355, 2080, 1422), F_SMALL, NAVY)
    return save(image, "07_diagrama_entidad_relacion")


def set_run_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0].cells
    for i, text in enumerate(headers):
        shade_cell(header[i], "E8EEF5")
        p = header[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=9.5, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            set_run_font(run, size=9.5, color=INK)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_paragraph(doc, text="", bold_prefix=None, style=None, align=None, after=6):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, color=INK, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=INK)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r, color=INK)


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, "F3F6F8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    for line_no, line in enumerate(code.strip().splitlines()):
        if line_no:
            p.add_run("\n")
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=8.5, color="22313F")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(path), width=Inches(6.25))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    run = cap.add_run(caption)
    set_run_font(run, size=9.5, color=MUTED, italic=True)


def add_screenshot(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(path), width=Inches(6.25))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    run = cap.add_run(caption)
    set_run_font(run, size=9.5, color=MUTED, italic=True)


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("ETRAL | Documentacion tecnica")
    set_run_font(run, size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Tesis de Ingenieria Industrial | Sistema web y gemelo digital")
    set_run_font(run, size=8.5, color=MUTED)


def build_doc(figures):
    doc = Document()
    configure_doc(doc)

    # Portada
    doc.add_paragraph().paragraph_format.space_after = Pt(82)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ETRAL")
    set_run_font(r, size=28, color=BLUE, bold=True)
    p.paragraph_format.space_after = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DOCUMENTACION TECNICA DEL SISTEMA WEB\nY GEMELO DIGITAL")
    set_run_font(r, size=22, color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Arquitectura, componentes, flujos, modelo de datos y evidencia de implementacion")
    set_run_font(r, size=12.5, color=MUTED, italic=True)
    p.paragraph_format.space_after = Pt(45)
    add_table(doc, ["Documento", "Uso"], [
        ("Anexo tecnico de tesis", "Sustentacion y entrega al jurado"),
        ("Proyecto", "Control de produccion y gemelo digital ETRAL"),
        ("Contenido", "Arquitectura, diagramas, codigo y trazabilidad"),
    ], [2700, 6660])
    doc.add_paragraph().paragraph_format.space_after = Pt(70)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Preparado para una tesis de Ingenieria Industrial")
    set_run_font(r, size=10, color=MUTED)
    doc.add_page_break()

    doc.add_heading("Contenido", level=1)
    add_bullets(doc, [
        "1. Proposito y alcance del documento",
        "2. Arquitectura general y diagrama de aplicaciones",
        "3. Componentes de la aplicacion web",
        "4. Flujos operativos y ejecucion del gemelo digital",
        "5. Modelo de datos y diagrama entidad-relacion",
        "6. Ejemplos de codigo y decisiones de diseno",
        "7. Evidencias y guion para sustentacion",
    ])
    doc.add_heading("1. Proposito y alcance", level=1)
    add_paragraph(doc, "Este documento explica la construccion tecnica del sistema ETRAL. El sistema registra productos, rutas, ordenes CECO, inventario, avances, recursos e incidencias de una planta. Sobre esta base opera un gemelo digital que permite comparar escenarios de capacidad, disponibilidad de materiales, demanda y mano de obra antes de ejecutar cambios en la operacion real.")
    add_paragraph(doc, "La finalidad es entregar al jurado una vista verificable de la arquitectura, el modelo de datos, los componentes y los flujos que conectan el registro operativo con el analisis de escenarios.")
    doc.add_heading("2. Arquitectura general", level=1)
    add_paragraph(doc, "La arquitectura esta dividida en capas para separar la experiencia de usuario, la coordinacion de servicios, las reglas de negocio y la infraestructura. Esta separacion reduce el acoplamiento y permite evolucionar el motor del gemelo digital sin reescribir los modulos de operacion.")
    add_table(doc, ["Capa", "Responsabilidad", "Implementacion"], [
        ("Presentacion", "Tableros, formularios, navegacion y registro diario.", "React + Vite; App.jsx, vistas y drawers."),
        ("Aplicacion", "Orquestar persistencia y seleccionar el motor de simulacion.", "repository.js y twinApi.js."),
        ("Dominio", "Calcular capacidad, alertas, MRP, atrasos y KPIs.", "simulator.js y services.py."),
        ("Infraestructura", "Persistir datos y exponer servicios HTTP.", "Supabase/PostgreSQL y FastAPI."),
    ], [1500, 4050, 3810])
    add_figure(doc, figures["architecture"], "Figura 1. Arquitectura por capas de ETRAL.")
    doc.add_heading("2.1 Diagrama de aplicaciones", level=2)
    add_paragraph(doc, "El usuario trabaja desde el navegador. La aplicacion React consulta o registra informacion mediante la capa de repositorio y ejecuta el gemelo mediante un adaptador. El motor puede ejecutarse en el navegador o en Python, segun la configuracion del entorno.")
    add_figure(doc, figures["applications"], "Figura 2. Aplicaciones y canales de comunicacion del sistema.")

    doc.add_page_break()
    doc.add_heading("3. Componentes de la aplicacion web", level=1)
    add_paragraph(doc, "App.jsx concentra la composicion de la experiencia. Cada vista cubre una responsabilidad operativa y solicita datos a un repositorio en lugar de consultar directamente la base de datos. Esta decision facilita pruebas y permite alternar entre datos de demostracion y Supabase.")
    add_figure(doc, figures["components"], "Figura 3. Componentes de interfaz, servicios y fuentes de datos.")
    doc.add_heading("3.1 Modulos funcionales", level=2)
    add_bullets(doc, [
        "ProductionBoard: visualiza ordenes, carga y situacion de produccion.",
        "ProductList y ProductFlowDrawer: administran productos y sus rutas.",
        "StagesView: registra fases, actividades y el flujo de fabricacion.",
        "InventoryView y BomPanel: controlan materiales, stock y estructura BOM.",
        "ExecutionPanel: registra avance, movimientos y controles de ejecucion.",
        "ResourcesView: administra turnos, personal, equipos, calendario y asignaciones.",
    ])
    doc.add_heading("4. Flujos operativos", level=1)
    add_paragraph(doc, "Los flujos muestran como el sistema conserva la trazabilidad de una orden y como el gemelo digital transforma datos operativos en informacion de decision.")
    doc.add_heading("4.1 Flujo de una orden CECO", level=2)
    add_figure(doc, figures["ceco"], "Figura 4. Flujo de registro, avance y cierre de una orden CECO.")
    add_paragraph(doc, "El control inicia con el pedido y la creacion del CECO. Luego el sistema valida que exista una ruta definida y que los materiales requeridos esten disponibles. Durante la ejecucion se registra WIP, avance de actividades y asignacion de recursos hasta llegar al control de calidad y cierre.")

    doc.add_page_break()
    doc.add_heading("4.2 Flujo de ejecucion del gemelo digital", level=2)
    add_figure(doc, figures["twin"], "Figura 5. Flujo de datos y calculo del gemelo digital.")
    add_paragraph(doc, "El gemelo recibe un escenario y construye un snapshot de la operacion. Esto significa que el analisis usa una copia de los datos para calcular resultados. La corrida no modifica el estado real de las ordenes, el inventario ni los avances de planta.")
    doc.add_heading("5. Modelo de datos", level=1)
    add_paragraph(doc, "El modelo relacional se organiza en maestros, planeamiento y operacion, recursos y control, y analitica. Esta estructura permite relacionar una orden con su producto, ruta, materiales, recursos, eventos de ejecucion y resultados de simulacion.")
    add_figure(doc, figures["database"], "Figura 6. Agrupacion logica de las tablas de base de datos.")
    doc.add_heading("5.1 Entidades relevantes", level=2)
    add_table(doc, ["Entidad", "Proposito", "Relacion principal"], [
        ("body_types", "Catalogo de productos fabricados.", "Define ruta y BOM."),
        ("ceco_orders", "Ordenes de produccion.", "Vincula producto, fase, avance y fecha."),
        ("inventory_items", "Existencias, reservas y stock de seguridad.", "Abastece BOM y movimientos."),
        ("stage_inventory", "Trabajo en proceso por fase.", "Relaciona CECO con una etapa."),
        ("resource_assignments", "Horas asignadas de personal a una actividad.", "Relaciona recurso, CECO y actividad."),
        ("simulation_runs", "Parametros y resultados de una corrida.", "Conserva evidencia analitica en JSONB."),
    ], [1950, 3920, 3490])

    doc.add_page_break()
    doc.add_heading("5.2 Diagrama entidad-relacion", level=2)
    add_figure(doc, figures["er"], "Figura 7. Relaciones principales del modelo entidad-relacion.")
    add_paragraph(doc, "Una orden CECO enlaza el producto y la fase actual con sus avances, WIP y recursos. El BOM enlaza el producto con los materiales requeridos. Estas relaciones son la base de la trazabilidad y de la simulacion de restricciones.")
    doc.add_heading("6. Ejemplos de codigo y decisiones de diseno", level=1)
    doc.add_heading("6.1 Patron Repository", level=2)
    add_paragraph(doc, "El patron Repository evita que las vistas conozcan los detalles de Supabase o de los datos de prueba. La interfaz pide un repositorio y utiliza un contrato de operaciones comun.")
    add_code_block(doc, '''import { localRepository } from "./localRepository.js";
import { hasSupabaseConfig, supabaseRepository } from "./supabaseRepository.js";

export function getRepository() {
  if (import.meta.env.VITE_DATA_MODE === "mock") return localRepository;
  return hasSupabaseConfig() ? supabaseRepository : supabaseRepository;
}''')
    doc.add_heading("6.2 Snapshot y adaptador del gemelo", level=2)
    add_paragraph(doc, "El snapshot concentra los datos que necesita la simulacion. El adaptador decide si la corrida se resuelve con JavaScript en el navegador o con la API Python.")
    add_code_block(doc, '''export async function runTwinSimulation(dataset, draft) {
  if (twinEngine !== "python") {
    return runDigitalTwin(dataset, draft);
  }

  const response = await fetch(`${baseUrl}/api/v1/simulations`, {
    method: "POST",
    headers: { "Content-Type": "application/json" },
    body: JSON.stringify({ name: "Simulacion desde interfaz", input: {
      snapshot: snapshotFromDataset(dataset)
    } }),
  });
  return response.json();
}''')

    doc.add_page_break()
    doc.add_heading("6.3 API de dominio", level=2)
    add_paragraph(doc, "La API FastAPI valida la entrada y delega la regla al servicio de dominio. El endpoint no contiene la logica analitica; por eso el motor es mas sencillo de probar y mantener.")
    add_code_block(doc, '''app = FastAPI(title="ETRAL Digital Twin API", version="0.1.0")

@app.post("/api/v1/mrp/evaluate")
def mrp(snapshot: FactorySnapshot) -> dict:
    return evaluate_mrp(snapshot)

@app.post("/api/v1/simulations")
def run_simulation(run: SimulationRun) -> dict:
    return {"name": run.name, "result": simulate_comparison(run.input)}''')
    doc.add_heading("6.4 Decisiones de diseno que se deben defender", level=2)
    add_bullets(doc, [
        "Separacion por capas: la interfaz, las reglas y la persistencia evolucionan con menor impacto entre si.",
        "Repository: permite trabajar con Supabase o datos de demostracion sin modificar las vistas.",
        "Adapter: permite elegir motor JavaScript o Python desde configuracion.",
        "Snapshot: protege la operacion real durante las pruebas de escenarios.",
        "DTO y schemas: validan los datos que viajan por la API del gemelo digital.",
    ])
    doc.add_heading("7. Evidencias y guion para sustentacion", level=1)
    add_paragraph(doc, "Para demostrar el funcionamiento, se recomienda presentar una orden CECO desde su registro hasta su cierre y luego ejecutar dos escenarios del gemelo digital: uno base y uno con una restriccion de demanda, inventario o capacidad.")
    add_table(doc, ["Evidencia", "Que demuestra"], [
        ("Captura de un CECO en dos o mas fases", "Trazabilidad del flujo operativo."),
        ("Consulta de ceco_orders, stage_inventory y ceco_activity_progress", "Persistencia y avance de la orden."),
        ("Movimiento de inventario y BOM asociado", "Consumo y disponibilidad de materiales."),
        ("Corrida base vs. corrida con restriccion", "Uso del gemelo para toma de decisiones."),
        ("Resultado almacenado en simulation_runs", "Evidencia y comparabilidad de escenarios."),
    ], [3000, 6360])
    doc.add_heading("7.1 Mensaje de cierre sugerido", level=2)
    add_paragraph(doc, "ETRAL integra el registro operacional con un gemelo digital orientado a decisiones. La base de datos conserva lo ocurrido en planta y el motor evalua escenarios sobre snapshots; por ello es posible anticipar restricciones sin poner en riesgo la operacion real.")
    doc.add_heading("Referencias de implementacion", level=1)
    add_bullets(doc, [
        "src/App.jsx y componentes de interfaz: composicion de la aplicacion.",
        "src/services/repository.js y src/services/supabaseRepository.js: acceso a datos.",
        "src/services/twinApi.js y src/lib/simulator.js: adaptador y motor JavaScript.",
        "backend/app/main.py y backend/app/services.py: API y motor Python.",
        "src/supabase/schema.sql: definicion de tablas y relaciones.",
    ])
    screenshots = [
        ("01_inicio.png", "Figura 8. Interfaz de inicio: indicadores de ordenes, bloqueos, materiales en riesgo y horas reportadas."),
        ("02_produccion.png", "Figura 9. Interfaz de produccion: tablero Kanban con ordenes CECO distribuidas por fase."),
        ("03_productos_rutas.png", "Figura 10. Interfaz de productos: plantillas maestras, rutas y componentes BOM."),
        ("04_fases_actividades.png", "Figura 11. Interfaz de fases y actividades: tiempos estandar, actividades e inventario en proceso."),
        ("05_inventario_bom.png", "Figura 12. Interfaz de inventario: existencias, comprometido, proyeccion y estado de cobertura."),
        ("06_recursos.png", "Figura 13. Interfaz de recursos: personal, turnos, equipos, asignaciones, calendario e incidencias."),
        ("07_simulacion.png", "Figura 14. Interfaz de simulacion: calibracion, horizonte de planificacion y configuracion del escenario."),
    ]
    available_screenshots = [(SCREENSHOT_DIR / name, caption) for name, caption in screenshots if (SCREENSHOT_DIR / name).exists()]
    if available_screenshots:
        doc.add_page_break()
        doc.add_heading("8. Capturas de las interfaces del sistema", level=1)
        add_paragraph(doc, "Las siguientes capturas documentan la interfaz real utilizada para registrar y consultar la operacion. Se incluyen como evidencia visual del alcance implementado y como apoyo para explicar la demostracion ante el jurado.")
        for index, (path, caption) in enumerate(available_screenshots):
            if index > 0:
                doc.add_page_break()
            add_screenshot(doc, path, caption)
    doc.save(OUTPUT)


def main():
    ASSETS.mkdir(parents=True, exist_ok=True)
    figures = {
        "architecture": create_architecture(),
        "applications": create_application_diagram(),
        "components": create_components(),
        "ceco": create_ceco_flow(),
        "twin": create_twin_flow(),
        "database": create_database_diagram(),
        "er": create_er_diagram(),
    }
    build_doc(figures)
    print(OUTPUT)


if __name__ == "__main__":
    main()
